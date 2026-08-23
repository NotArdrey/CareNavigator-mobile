from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
ASSETS = Path(__file__).resolve().parent
WEEK5 = ROOT / "CareNavigator_PH_Week_5_API_Integration_Submission.docx"
WEEK4 = ROOT / "CareNavigator_PH_Week_4_CRUD_Submission.docx"


def find_caption(doc: Document, prefix: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise ValueError(f"Caption not found: {prefix}")


def set_caption_text(paragraph, text: str):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def picture_paragraph_after(doc: Document, caption):
    paragraphs = doc.paragraphs
    start = paragraphs.index(caption)
    for paragraph in paragraphs[start + 1 :]:
        if paragraph._p.xpath(".//w:drawing"):
            return paragraph
        if paragraph.text.strip():
            break
    raise ValueError(f"Picture paragraph not found after: {caption.text}")


def replace_figure(
    doc: Document,
    caption_prefix: str,
    caption_text: str,
    image_path: Path,
    width_inches: float,
):
    caption = find_caption(doc, caption_prefix)
    set_caption_text(caption, caption_text)
    picture_paragraph = picture_paragraph_after(doc, caption)

    for run in list(picture_paragraph.runs):
        if run._r.xpath(".//w:drawing"):
            picture_paragraph._p.remove(run._r)

    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = picture_paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    inline = run._r.xpath(".//wp:inline")
    if inline:
        doc_pr = inline[0].find(qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", caption_text)


def update_week5():
    doc = Document(WEEK5)
    replace_figure(
        doc,
        "Figure 5.",
        "Figure 5. Mobile Flutter interface displaying hospitals retrieved from the live Supabase API.",
        ASSETS / "live_app_home_mobile.png",
        2.05,
    )
    doc.core_properties.comments = (
        "Updated with a 390 x 844 mobile-view screenshot for the live API evidence."
    )
    doc.save(WEEK5)


def update_week4():
    doc = Document(WEEK4)
    figures = [
        (
            "Figure 1.",
            "Figure 1. Mobile Add Slot form for creating a doctor schedule record.",
            "create_operation.png",
        ),
        (
            "Figure 2.",
            "Figure 2. Mobile doctor workspace displaying a persisted schedule record.",
            "read_operation.png",
        ),
        (
            "Figure 3.",
            "Figure 3. Mobile confirmation before publishing an inactive schedule slot.",
            "update_operation.png",
        ),
        (
            "Figure 4.",
            "Figure 4. Mobile destructive-action confirmation for deleting availability.",
            "delete_operation.png",
        ),
    ]
    for prefix, caption, filename in figures:
        replace_figure(doc, prefix, caption, ASSETS / filename, 2.35)
    doc.core_properties.comments = (
        "Updated with 390 x 844 mobile-view screenshots for all CRUD operations."
    )
    doc.save(WEEK4)


if __name__ == "__main__":
    update_week5()
    update_week4()
    print(WEEK5)
    print(WEEK4)
