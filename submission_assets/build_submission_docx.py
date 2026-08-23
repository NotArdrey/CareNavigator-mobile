from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = Path(__file__).resolve().parent
OUTPUT = ROOT / "CareNavigator_PH_API_CRUD_Submission.docx"

# compact_reference_guide preset, with a named editorial-cover title override.
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
TEAL = "0F766E"
TEAL_LIGHT = "E7F6F4"
PALE_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "64748B"
INK = "102A3A"
WHITE = "FFFFFF"
RED = "B4232C"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_shading(paragraph, fill: str, border: str = "D7E3EA"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "6")
        element.set(qn("w:color"), border)
        borders.append(element)
    p_pr.append(borders)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_begin, instr, fld_sep, text, fld_end))
    set_run_font(run, size=9, color=MID_GRAY)


def create_numbering(doc: Document):
    numbering = doc.part.numbering_part.element

    def add_abstract(abstract_id, fmt, text, left, hanging):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        level.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        level.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        level.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        level.append(p_pr)
        abstract.append(level)
        numbering.append(abstract)

    add_abstract(990, "bullet", "•", 540, 270)
    add_abstract(991, "decimal", "%1.", 540, 270)
    for num_id, abstract_id in ((990, 990), (991, 991)):
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_num_id = OxmlElement("w:abstractNumId")
        abstract_num_id.set(qn("w:val"), str(abstract_id))
        num.append(abstract_num_id)
        numbering.append(num)
    return 990, 991


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_id_el))
    p_pr.append(num_pr)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = rgb(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = rgb(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

caption = styles["Caption"]
caption.font.name = "Calibri"
caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
caption.font.size = Pt(9)
caption.font.italic = True
caption.font.color.rgb = rgb(MID_GRAY)
caption.paragraph_format.space_before = Pt(4)
caption.paragraph_format.space_after = Pt(4)
caption.paragraph_format.keep_with_next = True

bullet_num_id, decimal_num_id = create_numbering(doc)

header = section.header
header_p = header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_p.paragraph_format.space_after = Pt(0)
hr = header_p.add_run("CareNavigator PH  |  API + CRUD Submission")
set_run_font(hr, size=9, color=MID_GRAY, bold=True)

footer = section.footer
footer_p = footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = footer_p.add_run("CareNavigator PH  •  Page ")
set_run_font(fr, size=9, color=MID_GRAY)
add_page_number(footer_p)


def add_body(text, bold_label=None):
    p = doc.add_paragraph()
    if bold_label and text.startswith(bold_label):
        label = p.add_run(bold_label)
        set_run_font(label, bold=True)
        rest = p.add_run(text[len(bold_label):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    apply_num(p, bullet_num_id)
    set_run_font(p.add_run(text))
    return p


def add_callout(label, text, fill=TEAL_LIGHT, border="A7DAD4"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    set_paragraph_shading(p, fill, border)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, color=NAVY)
    set_run_font(p.add_run(text), color=INK)
    return p


def add_mono(text, size=9.5, color=DARK_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, name="Cascadia Mono", size=size, color=color)
    return p


def add_figure(path: Path, caption_text: str, width=6.5, max_height=6.4):
    cap = doc.add_paragraph(style="Caption")
    set_run_font(cap.add_run(caption_text), size=9, color=MID_GRAY, italic=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    inline = run._r.xpath(".//wp:inline")
    if inline:
        doc_pr = inline[0].find(qn("wp:docPr"))
        if doc_pr is not None:
            doc_pr.set("descr", caption_text)
    return p


def add_page_title(kicker, title, subtitle=None):
    k = doc.add_paragraph()
    k.paragraph_format.space_before = Pt(0)
    k.paragraph_format.space_after = Pt(4)
    kr = k.add_run(kicker.upper())
    set_run_font(kr, size=9.5, color=TEAL, bold=True)
    t = doc.add_paragraph()
    t.paragraph_format.space_before = Pt(0)
    t.paragraph_format.space_after = Pt(6)
    tr = t.add_run(title)
    set_run_font(tr, size=24, color=NAVY, bold=True)
    if subtitle:
        s = doc.add_paragraph()
        s.paragraph_format.space_before = Pt(0)
        s.paragraph_format.space_after = Pt(12)
        sr = s.add_run(subtitle)
        set_run_font(sr, size=12.5, color=MID_GRAY)


def page_break():
    doc.add_page_break()


# Cover / editorial_cover pattern.
spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(72)
add_page_title(
    "Course submission • Due August 17, 2026",
    "CareNavigator PH",
    "API Integration & HTTP Requests (Week 5)\nCRUD Functionality (Week 4)",
)
add_callout(
    "Document order",
    "Week 5 is presented first. Week 4 begins after a dedicated divider page so the two submission sets remain clearly separated.",
)
add_body("Project type: Flutter web/mobile healthcare navigation application", "Project type:")
add_body("Backend: Supabase Data API, PostgreSQL, Auth, Realtime, Storage, and Edge Functions", "Backend:")
add_body("Prepared: August 15, 2026", "Prepared:")
add_body("Student: ______________________________________________", "Student:")
add_body("Course / Section: _____________________________________", "Course / Section:")

doc.add_paragraph().paragraph_format.space_after = Pt(14)
for item in (
    "Live GET evidence from the configured Supabase project",
    "HTTP POST implementation with validation and JSON processing",
    "Database-backed Create, Read, Update, and Delete for doctor schedule slots",
    "Reproducible UI screenshots and passing verification results",
):
    add_bullet(item)


page_break()
add_page_title(
    "Part I • Week 5",
    "API Integration & HTTP Requests",
    "External/custom backend API integration for dynamic care-directory data and consultation workflows.",
)
add_callout(
    "API name",
    "Supabase Data API (PostgREST) with the CareNavigator Edge Function API. The integration retrieves verified hospital data, creates consultation requests, and supplies the care assistant with structured JSON.",
)
doc.add_paragraph("API endpoints and methods", style="Heading 2")
endpoint_rows = [
    ("GET", "/rest/v1/hospitals", "Retrieves verified, open or limited hospitals for the public directory."),
    ("POST", "/rest/v1/guest_consultation_requests", "Creates a validated consultation request and returns the new record ID."),
    ("POST", "/functions/v1/care-navigator-chat", "Sends message/facility JSON to the care-assistant Edge Function."),
]
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.rows[0].cells[0].text = "Method"
table.rows[0].cells[1].text = "Endpoint"
table.rows[0].cells[2].text = "Purpose"
for cell in table.rows[0].cells:
    shade_cell(cell, PALE_BLUE)
    for run in cell.paragraphs[0].runs:
        set_run_font(run, size=9.5, color=NAVY, bold=True)
mark_header_row(table.rows[0])
for method, endpoint, purpose in endpoint_rows:
    row = table.add_row()
    for index, value in enumerate((method, endpoint, purpose)):
        row.cells[index].text = value
        for run in row.cells[index].paragraphs[0].runs:
            set_run_font(run, name="Cascadia Mono" if index == 1 else "Calibri", size=9)
        if index == 0:
            row.cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
set_table_geometry(table, [1150, 3700, 4510])

doc.add_paragraph("Features added through the integration", style="Heading 2")
for item in (
    "A live public directory of verified Philippine hospitals, including operating state and availability information.",
    "Dynamic display of API data in Flutter rather than hard-coded records.",
    "Validated consultation-request creation through a POST request.",
    "Care-assistant JSON exchange through a protected Supabase Edge Function.",
    "Explicit loading, empty, unavailable, retry, and malformed-response handling.",
):
    add_bullet(item)

page_break()
add_page_title("Week 5 evidence 1", "HTTP GET request implementation")
add_figure(
    ASSETS / "api_get_implementation.png",
    "Figure 1. Supabase query used to retrieve verified, operating hospitals.",
)
add_body(
    "The Supabase Dart client composes an HTTP GET request to the hospitals REST endpoint. The query selects required JSON fields, filters verification_status to verified, restricts operating_status to open or limited, and orders results by hospital_name."
)
add_callout(
    "Meaningful processing",
    "The repository immediately returns an empty list for an empty response, then joins related departments, services, emergency-room records, bed capacity, doctors, and schedules before creating typed UI models.",
    fill=LIGHT_GRAY,
    border="D8E0E7",
)


page_break()
add_page_title("Week 5 evidence 2", "HTTP POST request implementation")
add_figure(
    ASSETS / "api_post_implementation.png",
    "Figure 2. POST-style insert for a validated guest consultation request.",
)
add_body(
    "The repository sends a JSON request body to guest_consultation_requests, requests the inserted id, parses the returned object, and exposes the ID to the calling UI. Authentication and field validation run before the request."
)
add_callout(
    "Validation",
    "Required fields, email format, concern length, hospital/department selection, supported consultation type, future schedule time, and verified-email ownership are checked before data is sent.",
    fill=LIGHT_GRAY,
    border="D8E0E7",
)


page_break()
add_page_title("Week 5 evidence 3", "JSON parsing and processing")
add_figure(
    ASSETS / "json_processing_implementation.png",
    "Figure 3. JSON rows converted into typed HospitalDirectoryEntry models.",
)
add_body(
    "Fields are read defensively with null-safe conversions. Operating and emergency states are combined into a meaningful isAvailable value, while geographic coordinates and capacity fields are normalized before the UI receives them."
)
add_callout(
    "Invalid responses",
    "PostgREST, Edge Function, authentication, repository, and unexpected exceptions are converted into user-safe messages. Raw errors and credentials are not displayed to end users.",
    fill=LIGHT_GRAY,
    border="D8E0E7",
)


page_break()
add_page_title("Week 5 evidence 4", "Live JSON response sample")
add_figure(
    ASSETS / "api_json_response.png",
    "Figure 4. Live GET response captured from the configured Supabase hospitals endpoint on August 15, 2026.",
    width=5.85,
)
add_callout(
    "Result",
    "The response returned verified, open facilities. The first two records correspond to the facilities rendered in the application screenshot on the next page.",
)
add_body(
    "Security note: the response sample intentionally excludes authentication headers and publishable-key values. Row Level Security remains the authorization boundary for Supabase data access."
)


page_break()
add_page_title("Week 5 evidence 5", "Retrieved API data displayed in the UI")
add_figure(
    ASSETS / "live_app_home_mobile.png",
    "Figure 5. Mobile Flutter interface displaying hospitals retrieved from the live Supabase API.",
    width=2.05,
)
doc.add_paragraph("Loading and response-state handling", style="Heading 2")
for item in (
    "Loading: the panel shows “Loading verified facilities” while the repository request is active.",
    "Error: a “Hospital directory unavailable” state displays a safe message and Retry button.",
    "Empty: “No facilities published” appears when the API returns no matching records.",
    "Success: verified facilities are rendered with name, location, operating state, directions, bed availability, and clinician count.",
):
    add_bullet(item)
add_callout(
    "Verification",
    "The live GET response and the rendered UI were captured on August 15, 2026. flutter analyze completed with no issues.",
    fill=LIGHT_GRAY,
    border="D8E0E7",
)


# Full-page divider: the requested lower section begins here.
page_break()
spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(54)
add_page_title(
    "Part II • Week 4",
    "Implement CRUD Functionality",
    "Complete Create, Read, Update, and Delete operations using persistent Supabase PostgreSQL storage.",
)
add_callout(
    "Main data entity",
    "Doctor schedule slots (doctor_schedules). Each record stores the doctor, weekday, start/end time, consultation type, appointment duration, and active state.",
)
doc.add_paragraph("CRUD implementation summary", style="Heading 2")
crud_rows = [
    ("Create", "createScheduleSlot", "INSERT / POST", "Validates day, time order, duration, and care mode."),
    ("Read", "WorkspaceRepository.load", "SELECT / GET", "Filters rows to the authenticated doctor."),
    ("Update", "setScheduleActive", "UPDATE / PATCH", "Requires a non-empty schedule ID."),
    ("Delete", "deleteScheduleSlot", "DELETE", "Blocks removal when an appointment depends on the slot."),
]
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
for index, value in enumerate(("Operation", "Repository method", "Action", "Safeguard")):
    table.rows[0].cells[index].text = value
    shade_cell(table.rows[0].cells[index], PALE_BLUE)
    for run in table.rows[0].cells[index].paragraphs[0].runs:
        set_run_font(run, size=9, color=NAVY, bold=True)
mark_header_row(table.rows[0])
for values in crud_rows:
    row = table.add_row()
    for index, value in enumerate(values):
        row.cells[index].text = value
        for run in row.cells[index].paragraphs[0].runs:
            set_run_font(run, name="Cascadia Mono" if index == 1 else "Calibri", size=8.8)
        if index in (0, 2):
            row.cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
set_table_geometry(table, [1150, 2700, 1450, 4060])

doc.add_paragraph("Persistence and user experience", style="Heading 2")
for item in (
    "Records are persisted in Supabase PostgreSQL and protected by authentication and Row Level Security.",
    "The UI provides labeled forms, status chips, publish/unpublish actions, confirmation dialogs, and retry states.",
    "Destructive deletion is guarded by a booked-appointment dependency check.",
):
    add_bullet(item)


def crud_page(number, operation, image_name, caption_text, description, method_text, callout_label, callout_text, callout_fill=TEAL_LIGHT, callout_border="A7DAD4"):
    page_break()
    add_page_title(f"Week 4 evidence {number}", f"{operation} operation")
    add_figure(ASSETS / image_name, caption_text, width=2.35)
    add_body(description)
    add_mono(method_text)
    add_callout(callout_label, callout_text, fill=callout_fill, border=callout_border)


crud_page(
    1,
    "Create",
    "create_operation.png",
    "Figure 6. Add Slot form for creating a doctor schedule record.",
    "The Create form collects a weekday, 24-hour start/end times, care mode, and minutes per appointment. Submitting calls the repository insert and refreshes the persisted schedule list.",
    "CareRepository.createScheduleSlot(...) → doctor_schedules.insert({...})",
    "Validation",
    "Day must be 0-6, times must use HH:mm, end must be later than start, duration must be 10-240 minutes, and consultation type is normalized.",
)

crud_page(
    2,
    "Read",
    "read_operation.png",
    "Figure 7. Persisted schedule record displayed in the doctor workspace.",
    "The Read flow selects doctor_schedules through WorkspaceRepository, limits records to the authenticated doctor, orders them, maps JSON rows to WorkspaceItem models, and annotates booking dependencies.",
    "WorkspaceRepository.load(...) → from('doctor_schedules').select(...)",
    "Meaningful display",
    "The UI shows the weekday, time range, care mode, slot duration, active state, and available row actions instead of raw JSON.",
)

crud_page(
    3,
    "Update",
    "update_operation.png",
    "Figure 8. Confirmation before publishing an inactive schedule slot.",
    "The Update flow toggles is_active for the selected schedule ID. A confirmation dialog explains that publishing changes whether patients can see and use the recurring availability.",
    "CareRepository.setScheduleActive(...) → update({'is_active': active})",
    "Safety",
    "The repository rejects an empty schedule ID, updates only the matching row, and requests the updated ID so a missing/invalid record does not silently succeed.",
)

crud_page(
    4,
    "Delete",
    "delete_operation.png",
    "Figure 9. Destructive-action confirmation for deleting availability.",
    "The Delete flow first confirms ownership, compares the recurring slot with non-cancelled appointments, and removes the record only when no booking depends on it.",
    "CareRepository.deleteScheduleSlot(...) → delete().eq('id', scheduleId)",
    "Protected deletion",
    "If a booked consultation matches the slot, deletion is blocked with a clear message. The final confirmation uses explicit destructive wording.",
    callout_fill="FDEBEC",
    callout_border="E8A0A5",
)


page_break()
add_page_title("Week 4 documentation", "Validation, persistence, and verification")
doc.add_paragraph("How CRUD was implemented", style="Heading 2")
add_body(
    "The Flutter UI delegates all database work to typed repository classes. Supabase translates select, insert, update, and delete calls into HTTP requests to PostgreSQL through PostgREST. Riverpod reloads the workspace snapshot after successful mutations, so users immediately see persisted state."
)
doc.add_paragraph("Data validation", style="Heading 2")
for item in (
    "Client validation prevents malformed day, time, duration, care-mode, and identifier values before a network request.",
    "Database constraints and RLS provide server-side persistence and access control.",
    "Delete checks protect appointments that depend on a recurring slot.",
    "Repository failures are converted into concise messages suitable for the UI.",
):
    add_bullet(item)

doc.add_paragraph("Verification results", style="Heading 2")
add_callout(
    "Static analysis",
    "flutter analyze: No issues found (August 15, 2026).",
    fill=LIGHT_GRAY,
    border="D8E0E7",
)
add_callout(
    "Focused tests",
    "30 repository-validation, primary-action, row-action, and CRUD-evidence tests passed. The four screenshots in this section were generated by a reproducible Flutter widget test without modifying live medical data.",
)

doc.add_paragraph("Submission checklist", style="Heading 2")
for item in (
    "Updated project source code: present in the CareNavigator PH workspace.",
    "API request implementation screenshots: Figures 1 and 2.",
    "Retrieved API data screenshot: Figure 5.",
    "JSON response sample: Figure 4.",
    "Create, Read, Update, and Delete screenshots: Figures 6-9.",
    "Brief API and CRUD documentation: included in Parts I and II.",
):
    add_bullet(item)

add_body("End of submission document.")

doc.core_properties.title = "CareNavigator PH API Integration and CRUD Submission"
doc.core_properties.subject = "Week 5 API Integration and Week 4 CRUD Functionality"
doc.core_properties.author = "CareNavigator PH Student Submission"
doc.core_properties.keywords = "Flutter, Supabase, API, HTTP, JSON, CRUD"
doc.core_properties.comments = "Generated from the verified CareNavigator PH project workspace."

doc.save(OUTPUT)
print(OUTPUT)
