from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "CareNavigator_PH_API_CRUD_Submission.docx"
WEEK5_OUTPUT = ROOT / "CareNavigator_PH_Week_5_API_Integration_Submission.docx"
WEEK4_OUTPUT = ROOT / "CareNavigator_PH_Week_4_CRUD_Submission.docx"

NAVY = "0B2545"
TEAL = "0F766E"
MID_GRAY = "64748B"
INK = "102A3A"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold


def element_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def is_page_break(element) -> bool:
    for br in element.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def body_children(doc: Document):
    return list(doc._element.body)


def marker_index(doc: Document, marker: str) -> int:
    for index, element in enumerate(body_children(doc)):
        if marker in element_text(element):
            return index
    raise ValueError(f"Marker not found: {marker}")


def preceding_page_break_index(doc: Document, marker: str) -> int:
    children = body_children(doc)
    index = marker_index(doc, marker)
    for candidate in range(index - 1, -1, -1):
        if is_page_break(children[candidate]):
            return candidate
    raise ValueError(f"No page break found before: {marker}")


def remove_body_range(doc: Document, start: int, end: int):
    body = doc._element.body
    children = body_children(doc)
    for element in children[start:end]:
        if element.tag != qn("w:sectPr"):
            body.remove(element)


def replace_in_runs(paragraph, old: str, new: str) -> bool:
    changed = False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            changed = True
    return changed


def find_paragraph(doc: Document, starts_with: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(starts_with):
            return paragraph
    raise ValueError(f"Paragraph not found: {starts_with}")


def set_paragraph_run_text(paragraph, text: str):
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def apply_num(paragraph, num_id=990):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_id_el))
    p_pr.append(num_pr)


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(27)
    paragraph.paragraph_format.first_line_indent = Pt(-13.5)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    apply_num(paragraph)
    set_run_font(paragraph.add_run(text))


def add_page_title(doc: Document, kicker: str, title: str, subtitle: str | None = None):
    kicker_p = doc.add_paragraph()
    kicker_p.paragraph_format.space_after = Pt(4)
    set_run_font(kicker_p.add_run(kicker.upper()), size=9.5, color=TEAL, bold=True)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(6)
    set_run_font(title_p.add_run(title), size=24, color=NAVY, bold=True)

    if subtitle:
        subtitle_p = doc.add_paragraph()
        subtitle_p.paragraph_format.space_after = Pt(12)
        set_run_font(subtitle_p.add_run(subtitle), size=12.5, color=MID_GRAY)


def update_header(doc: Document, text: str):
    paragraph = doc.sections[0].header.paragraphs[0]
    set_paragraph_run_text(paragraph, text)
    if paragraph.runs:
        set_run_font(paragraph.runs[0], size=9, color=MID_GRAY, bold=True)


def update_cover(doc: Document, subtitle: str, scope: str, bullets: list[str]):
    cover_subtitle = find_paragraph(doc, "API Integration & HTTP Requests (Week 5)")
    set_paragraph_run_text(cover_subtitle, subtitle)

    scope_paragraph = find_paragraph(doc, "Document order:")
    if scope_paragraph.runs:
        scope_paragraph.runs[0].text = "Submission scope: "
        if len(scope_paragraph.runs) > 1:
            scope_paragraph.runs[1].text = scope
        else:
            scope_paragraph.add_run(scope)

    original_bullets = [
        "Live GET evidence from the configured Supabase project",
        "HTTP POST implementation with validation and JSON processing",
        "Database-backed Create, Read, Update, and Delete for doctor schedule slots",
        "Reproducible UI screenshots and passing verification results",
    ]
    for old, new in zip(original_bullets, bullets):
        set_paragraph_run_text(find_paragraph(doc, old), new)


def build_week5():
    doc = Document(SOURCE)
    week4_start = preceding_page_break_index(doc, "PART II")
    sect_index = next(
        index
        for index, element in enumerate(body_children(doc))
        if element.tag == qn("w:sectPr")
    )
    remove_body_range(doc, week4_start, sect_index)

    update_header(doc, "CareNavigator PH  |  Week 5 API Integration Submission")
    update_cover(
        doc,
        "API Integration & HTTP Requests (Week 5)",
        "This file contains only the Week 5 API Integration and HTTP Requests submission requirements and evidence.",
        [
            "Live GET evidence from the configured Supabase project",
            "HTTP POST implementation with validation and JSON processing",
            "Retrieved API data displayed in the Flutter user interface",
            "Loading, invalid-response, empty-state, and success-state handling",
        ],
    )
    set_paragraph_run_text(find_paragraph(doc, "PART I"), "WEEK 5 SUBMISSION")

    doc.add_page_break()
    add_page_title(doc, "Week 5 documentation", "Submission checklist")
    for item in (
        "Updated project source code: present in the CareNavigator PH workspace.",
        "API request implementation screenshots: Figures 1 and 2.",
        "JSON parsing implementation screenshot: Figure 3.",
        "Live JSON response sample: Figure 4.",
        "Retrieved API data displayed in the application: Figure 5.",
        "API name, endpoints, request methods, features, loading states, and invalid-response handling are documented.",
    ):
        add_bullet(doc, item)
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(end_p.add_run("End of Week 5 submission document."))

    doc.core_properties.title = "CareNavigator PH Week 5 API Integration Submission"
    doc.core_properties.subject = "API Integration and HTTP Requests - Week 5"
    doc.core_properties.keywords = "Flutter, Supabase, API, HTTP, GET, POST, JSON"
    doc.save(WEEK5_OUTPUT)


def build_week4():
    doc = Document(SOURCE)
    week5_start = preceding_page_break_index(doc, "PART I")
    week4_start = preceding_page_break_index(doc, "PART II")
    remove_body_range(doc, week5_start, week4_start)

    update_header(doc, "CareNavigator PH  |  Week 4 CRUD Submission")
    update_cover(
        doc,
        "Implement CRUD Functionality (Week 4)",
        "This file contains only the Week 4 Create, Read, Update, and Delete submission requirements and evidence.",
        [
            "Database-backed Create, Read, Update, and Delete for doctor schedule slots",
            "Client validation and protected destructive actions",
            "Persistent Supabase PostgreSQL storage with Row Level Security",
            "Create, Read, Update, and Delete user-interface screenshots",
        ],
    )
    set_paragraph_run_text(find_paragraph(doc, "PART II"), "WEEK 4 SUBMISSION")

    caption_replacements = {
        "Figure 6.": "Figure 1.",
        "Figure 7.": "Figure 2.",
        "Figure 8.": "Figure 3.",
        "Figure 9.": "Figure 4.",
    }
    for paragraph in doc.paragraphs:
        for old, new in caption_replacements.items():
            replace_in_runs(paragraph, old, new)

    checklist_items = [
        "Updated project source code: present in the CareNavigator PH workspace.",
        "API request implementation screenshots: Figures 1 and 2.",
        "Retrieved API data screenshot: Figure 5.",
        "JSON response sample: Figure 4.",
        "Create, Read, Update, and Delete screenshots: Figures 6-9.",
        "Brief API and CRUD documentation: included in Parts I and II.",
    ]
    replacement_items = [
        "Updated project source code: present in the CareNavigator PH workspace.",
        "Create operation screenshot: Figure 1.",
        "Read operation screenshot: Figure 2.",
        "Update operation screenshot: Figure 3.",
        "Delete operation screenshot: Figure 4.",
        "Brief CRUD implementation, validation, persistence, and verification documentation is included.",
    ]
    for old, new in zip(checklist_items, replacement_items):
        set_paragraph_run_text(find_paragraph(doc, old), new)

    set_paragraph_run_text(
        find_paragraph(doc, "End of submission document."),
        "End of Week 4 submission document.",
    )
    doc.core_properties.title = "CareNavigator PH Week 4 CRUD Submission"
    doc.core_properties.subject = "CRUD Functionality - Week 4"
    doc.core_properties.keywords = "Flutter, Supabase, PostgreSQL, Create, Read, Update, Delete"
    doc.save(WEEK4_OUTPUT)


if __name__ == "__main__":
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    build_week5()
    build_week4()
    print(WEEK5_OUTPUT)
    print(WEEK4_OUTPUT)
